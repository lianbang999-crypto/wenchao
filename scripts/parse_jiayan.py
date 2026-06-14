# -*- coding: utf-8 -*-
"""
《印光法师嘉言录》白话（东林寺文库）docx → 实验模型 JSON

体例（已探查验证）：
  - 十编主题标题 = Normal 样式短行（一、赞净土超胜 …… 十、标应读典籍）
  - 子节 = 甲、乙、丙… 短行；无子节的编整编为一篇
  - 条目 = 【原文】段（可多段）+【译文】段（可多段），逐条对照
  - Body Text 样式行均为页眉页码（"编名 · 15 ·"、"· 2 书名 白话"、"目录 N"），过滤
  - 书首：印光法师语 / 封面题词 / 封二题词等小品；书尾：东林寺附录
铁律：只切分、过滤版式噪音，不改一字；逐篇对账（忽略空白）必须一致。

输出：build/jy/{index.json, articles/}，条目存为 {os:[原文段], ts:[译文段]}，
      与 migrate_v2.py 衔接（os/ts 等长则前端逐段对照，不等则分块）。
"""
import json
import os
import re

import docx

PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(PROJ, '印祖文钞', '嘉言录菁华录等',
                   '03印光法师嘉言录白话（东林寺文库）2023.11.26.docx')
OUT = os.path.join(PROJ, 'build', 'jy')
REPORT = os.path.join(PROJ, 'build', 'reports', 'report_jy.txt')

# 十编（原书目次，白名单精确匹配）
BIAN = ['一、赞净土超胜', '二、诫信愿真切', '三、示修持方法', '四、论生死事大',
        '五、勉居心诚敬', '六、告注重因果', '七、分禅净界限', '八、释普通疑惑',
        '九、谕在家善信', '十、标应读典籍']
RE_SUB = re.compile(r'^[甲乙丙丁戊己庚辛]、.{2,12}$')      # 子节
# 条目出处行：卷三·序·重刻阿弥陀经序（全书 434 条，结构化为 src）
RE_SRC = re.compile(r'^卷[一二三四]·[^·]{1,6}·.{2,24}$')
# 书尾锚点标题 → 所属分部
ANCHORS = {'修订说明': '附录', '东林佛号曲谱': '附录', '念佛的心态与音声': '附录', '回向偈': '附录'}
RE_JUNK = [
    re.compile(r'·\s*\d+\s*·'),          # 页眉：编名 · 15 ·
    re.compile(r'^·\s*\d+\s+'),          # 页眉：· 2  书名 白话
    re.compile(r'^\d+\s*·\s'),            # 页眉：224 ·  书名 白话
    re.compile(r'^目录\s*\d*$'),
    re.compile(r'^\d+\s*$'),              # 纯页码
    re.compile(r'^附\s*录\s*[.·]\s*\d+$'),  # 页眉：附 录 . 483
    re.compile(r'^增附\s*\d+$'),           # 页眉：增附 469
    re.compile(r'^0\d{2,3}-\d{6,8}$'),     # 版权页电话
    re.compile(r'^[\w.]+@[\w.]+$'),        # 版权页邮箱
    re.compile(r'^庐山东林寺官方平台$'),
    # 页眉页脚：整行仅「页码装饰 + 书名 + 白话」（含无"·"的页脚，如 "6  印光法师嘉言录 白话"）
    re.compile(r'^[·\d\s]*印光法师嘉言录\s*白话[·\d\s]*$'),
]


def is_junk(t, style):
    if any(r.search(t) for r in RE_JUNK):
        return True
    return False


# 边界字符：汉字 + 中文标点 + 全角符号 + 中英文弯引号（引号旁的排版空格也清除）
_CJKB = '一-鿿　-〿＀-￯‘’“”〝〞'
RE_CJK_GAP = re.compile('(?<=[' + _CJKB + '])[ \t\xa0　]+(?=[' + _CJKB + '])')
# 段首编号：把被物理换行拆散的整理类正文按编号项重新分段（含 1. / 一、/（一）/(1) 等）
RE_PARA_START = re.compile(
    r'^\s*(\d+\s*[.．、]|[①-⑩]|[一二三四五六七八九十]+\s*[、.]|[（(][一二三四五六七八九十\d]{1,3}[）)])')


def tidy(s):
    """清除汉字（含中文标点）之间的排版空格——东林本 PDF 转档遗留，
    非内容空白；对账忽略空白，不影响逐字一致性。"""
    return RE_CJK_GAP.sub('', s)


def merge_plain(lines):
    """把被 PDF 物理换行拆散的 plain 段落重新拼回完整段落。
    编号/序号行（1. / 一、/（一）)另起一段；纯标题短行（如"一、极乐世界的音声特质"，
    短且无尾标点）独立成段、其后正文另起；其余续接上一段。
    仅整理版式换行，不增删一字；对账走独立流水，不受影响。"""
    paras = []
    after_head = False     # 上一行是独立小标题 → 本行另起，不并入标题
    for ln in lines:
        head = bool(RE_PARA_START.match(ln)) and len(ln) <= 16 and not re.search(r'[。！？，、；：]$', ln)
        if not paras or RE_PARA_START.match(ln) or after_head:
            paras.append(ln)
        else:
            paras[-1] += ln
        after_head = head
    return paras


def parse_xuandu_list(lines):
    """把篇目行（保留空格）切成分组：按「（以上 X）」标记归组，◎ 标"尤要"。
    返回 [{'sec':'卷一·书一','items':[{'t':篇名,'m':是否◎}, …]}, …]"""
    sections, cur = [], []

    def add(seg):
        seg = seg.strip()
        if not seg:
            return
        marked = seg.startswith('◎')      # ◎ 只标紧跟的首篇（同行其余为另起篇目）
        seg = seg.lstrip('◎').strip()
        first = True
        for nm in re.split(r'\s+', seg):
            nm = nm.strip()
            if nm:
                cur.append({'t': nm, 'm': marked and first})
                first = False

    for line in lines:
        while True:
            mm = re.search(r'（以上\s*([^）]*)）', line)
            if not mm:
                break
            add(line[:mm.start()])
            sec = re.sub(r'\s+', '', mm.group(1)).replace('.', '·')
            sections.append({'sec': sec, 'items': cur})
            cur = []
            line = line[mm.end():]
        add(line)
    if cur:
        sections.append({'sec': '其他', 'items': cur})
    return sections


def build_xuandu(raw_lines):
    """从 docx 原始段落（保留空格）重建《印光法师文钞》选读：序文 + 分组篇目。
    返回 (序文str, sections)。"""
    start = next((i for i, t in enumerate(raw_lines)
                  if t.startswith('《印光法师文钞》选读')), None)
    if start is None:
        return '', []
    end = next((i for i in range(start + 1, len(raw_lines))
                if raw_lines[i] == '一、赞净土超胜'), len(raw_lines))
    preface, list_lines, in_list = [], [], False
    for i in range(start + 1, end):
        t = raw_lines[i]
        if not t or t.startswith('《印光法师文钞》选读') or is_junk(tidy(t), 'Normal'):
            continue
        if not in_list:
            preface.append(t)
            if '圆净谨志' in t:
                in_list = True
            continue
        list_lines.append(t)
    return tidy(''.join(preface)), parse_xuandu_list(list_lines)


def main():
    d = docx.Document(SRC)
    paras = [(tidy(p.text.strip()), p.style.name) for p in d.paragraphs]
    raw_lines = [p.text.strip() for p in d.paragraphs]   # 保留空格：选读篇目按空格切分用

    # 正文起点：「嘉言录封面题词」且其后 6 段内出现【原文】
    body_start = None
    for i, (t, st) in enumerate(paras):
        if t == '嘉言录封面题词' and any(
                paras[j][0].startswith('【原文】') for j in range(i + 1, min(i + 7, len(paras)))):
            body_start = i
            break
    assert body_start, '未找到正文起点'

    # 书首「印光法师语」（目录之前的题词页）
    epigraph = []
    for i, (t, st) in enumerate(paras[:body_start]):
        if t == '印光法师语':
            j = i + 1
            while j < body_start and not re.match(r'^目录', paras[j][0]):
                if paras[j][0] and not is_junk(*paras[j]):
                    epigraph.append(paras[j][0])
                j += 1
            break

    articles = []          # {title, part, plain:[前置段], entries:[{os,ts}]}
    report = [f'源文件: {os.path.basename(SRC)}', f'总段落: {len(paras)}  正文起点: {body_start}']
    dropped = []           # 被过滤的版式行（抽样记录）

    cur = None             # 当前文章
    part = '卷首'
    seen_bian = set()
    entry = None           # 当前条目 {'os':[], 'ts':[], src?}; side: 'o'|'t'
    side = 'o'
    stream = []            # 全局流水：按消费顺序记录每一行（对账用）
    skip_next = set()      # 拆行装饰头的后续行索引

    def flush_entry():
        nonlocal entry
        if entry and (entry['os'] or entry['ts']):
            cur['entries'].append(entry)
        entry = None

    def new_article(title, new_part=None):
        nonlocal cur, part, entry, side
        flush_entry()
        if cur and (cur['plain'] or cur['entries']):
            articles.append(cur)
        if new_part is not None:
            part = new_part
        cur = {'title': title, 'part': part, 'plain': [], 'entries': []}
        entry, side = None, 'o'

    def origin_follows(idx, limit=6):
        """其后 limit 个非空段内是否紧跟【原文】——用以判定真小品标题，
        与选读清单里以「序/记」结尾的文钞篇名（其后无【原文】）区分。"""
        seen = 0
        for j in range(idx + 1, len(paras)):
            tj, sj = paras[j]
            if not tj or is_junk(tj, sj):
                continue
            if tj in BIAN or tj.startswith('【译文】'):
                return False    # 已进入十编正文，此处【原文】不属于卷首小品
            if tj.startswith('【原文】'):
                return True
            seen += 1
            if seen >= limit:
                break
        return False

    if epigraph:
        cur = {'title': '印光法师语', 'part': '卷首', 'plain': epigraph, 'entries': []}
        articles.append(cur)
        cur = None

    for i in range(body_start, len(paras)):
        t, st = paras[i]
        if not t:
            continue
        if is_junk(t, st):
            if len(dropped) < 8:
                dropped.append(t[:24])
            continue
        # 同篇内重复出现的篇名行 = 分页题头，仅入流水、不入正文
        if cur is not None and t == cur['title']:
            stream.append(t)
            continue
        # 「附 录」分部行
        if re.sub(r'[\s　]+', '', t) == '附录' and len(seen_bian) == 10:
            part = '附录'
            stream.append(t)
            continue
        # 结构标题（编/子节不入流水，文档流同样排除）
        if t in BIAN:
            seen_bian.add(t)
            new_article(t, new_part=t)
            continue
        if RE_SUB.match(t) and st == 'Normal' and len(seen_bian) > 0 and len(seen_bian) < 10 + 1:
            # 子节：若当前编文章尚无内容，子节取代其成为文章
            if cur and cur['title'] == part and not cur['plain'] and not cur['entries']:
                cur['title'] = t
            else:
                new_article(t)
            continue
        # 「附 录」拆成两行的装饰头（"附"+"录"相邻单字行）：入流水、跳过
        if t == '附' and i + 2 < len(paras):
            nxt = next((j for j in range(i + 1, min(i + 3, len(paras))) if paras[j][0]), None)
            if nxt and paras[nxt][0] == '录':
                stream.append('附')
                part = '附录'
                skip_next.add(nxt)
                continue
        if i in skip_next:
            stream.append(t)
            continue
        stream.append(t)
        # 书尾锚点标题
        if t in ANCHORS:
            new_article(t, new_part=ANCHORS[t])
            continue
        # 原书「附」（编十之后的释疑附文）
        if t == '附' and len(seen_bian) == 10:
            new_article('附', new_part='增附')
            continue
        # 条目出处行（位于原文与译文之间）→ 结构化挂到当前条目，不冲洗
        if RE_SRC.match(t):
            if entry is not None:
                entry['src'] = t
            elif cur is not None:
                cur['plain'].append(t)
            continue
        # 条目标记
        if t.startswith('【原文】'):
            flush_entry()
            entry = {'os': [t[4:].strip()] if t[4:].strip() else [], 'ts': []}
            side = 'o'
            continue
        if t.startswith('【译文】'):
            if entry is None:
                entry = {'os': [], 'ts': []}
            side = 't'
            rest = t[4:].strip()
            if rest:
                entry['ts'].append(rest)
            continue
        # 卷首「《印光法师文钞》选读（附录于后）」：推荐篇目清单，整篇按原样列出
        # （否则其下数十条文钞篇名会被并入上一篇译文，且「…序/记」结尾的篇名会被误判为独立小品）
        if part == '卷首' and t.startswith('《印光法师文钞》选读'):
            new_article(t)
            cur['raw'] = True       # 篇目清单逐行保留，不做段落合并
            continue
        # 书首/书尾的小品标题：短行、无标点、具题名特征（会先冲洗未结条目）
        # 卷首小品须紧随【原文】，方与选读清单中的「…序/记」篇名区分开
        no_punct = len(t) <= 24 and st == 'Normal' and not re.search(r'[。，；：？！】]', t)
        is_front_title = (part == '卷首' and no_punct
                          and re.search(r'(题词|序|跋|语|记)$', t)
                          and origin_follows(i))
        # 附录小品标题须具题名特征；不再把「一、二、…」当独立篇——
        # 它们是附录文章（如「念佛的心态与音声」）内部的小节标题，应归入该篇
        is_appendix_title = (part in ('附录', '增附') and no_punct
                             and re.search(r'(后记|题词|序|跋)$', t))
        if is_front_title or is_appendix_title:
            new_article(t, new_part='附录' if is_appendix_title else None)
            continue
        # 普通内容行
        if cur is None:
            new_article('〔无题〕')
        if entry is not None:
            (entry['os'] if side == 'o' else entry['ts']).append(t)
        else:
            cur['plain'].append(t)
    new_article('〔收尾〕')  # 冲掉最后一篇（占位篇无内容不会入列）

    # ---- 对账：重建文本流 vs 文档原文（忽略空白与版式行）----
    def sq(x):
        return re.sub(r'\s+', '', x)

    doc_stream = []
    for i in range(body_start, len(paras)):
        t, st = paras[i]
        if not t or is_junk(t, st) or t in BIAN:
            continue
        if RE_SUB.match(t) and st == 'Normal':
            continue
        doc_stream.append(t)
    d_sq, j_sq = sq(''.join(doc_stream)), sq(''.join(stream))
    if d_sq != j_sq:
        pos = next((k for k in range(min(len(d_sq), len(j_sq))) if d_sq[k] != j_sq[k]),
                   min(len(d_sq), len(j_sq)))
        report.append(f'⚠ 对账不一致：文档{len(d_sq)}字 / JSON{len(j_sq)}字')
        report.append(f'  首分歧@{pos}: 文档[…{d_sq[max(0,pos-10):pos+14]}] JSON[…{j_sq[max(0,pos-10):pos+14]}]')
    else:
        report.append(f'全书对账一致 ✓（{len(d_sq)} 字）')

    # ---- 输出（实验模型）----
    os.makedirs(os.path.join(OUT, 'articles'), exist_ok=True)
    parts, idx = [], 0
    report.append(f'\n文章数: {len(articles)}  过滤版式行示例: {dropped}')
    report.append('\n--- 每篇概况 ---')
    for a in articles:
        idx += 1
        aid = f'{idx:03d}'
        segs = []
        xuandu = None
        if a['title'].startswith('《印光法师文钞》选读'):
            # 选读：序文单独成段，篇目清单结构化（分组 + ◎ + 待链接），不走 raw 列表
            preface, sections = build_xuandu(raw_lines)
            plain = [preface] if preface else []
            xuandu = sections
        else:
            plain = a['plain'] if a.get('raw') else merge_plain(a['plain'])
        for p in plain:
            segs.append({'o': p})
        for e in a['entries']:
            # 每条原文/译文各自被 PDF 物理换行拆成多段，在此各合回一整段，
            # 使原文与译文 1:1 等长 → 前端逐句对照（修复"没有对齐"）
            o = ''.join(e['os']).strip()
            t2 = ''.join(e['ts']).strip()
            seg = {'os': [o] if o else [], 'ts': [t2] if t2 else []}
            if e.get('src'):
                seg['src'] = e['src']
            segs.append(seg)
        art = {'id': aid, 'title': a['title'], 'translator': '', 'summary': '',
               'part': a['part'], 'segments': segs}
        if xuandu is not None:
            art['xuandu'] = xuandu
        with open(os.path.join(OUT, 'articles', aid + '.json'), 'w', encoding='utf-8') as f:
            json.dump(art, f, ensure_ascii=False, separators=(',', ':'))
        if not parts or parts[-1]['title'] != (a['part'] or '其他'):
            parts.append({'title': a['part'] or '其他', 'articles': []})
        parts[-1]['articles'].append({'id': aid, 'title': a['title'], 'summary': ''})
        n_pair = sum(1 for e in a['entries'] if ''.join(e['os']).strip() and ''.join(e['ts']).strip())
        report.append(f"{aid} {a['title'][:20]:<22} 条目{len(a['entries']):>3}（对照{n_pair}） 前置段{len(plain)}  [{a['part']}]")
    index = {'id': 'jy', 'title': '印光法师嘉言录（白话）', 'parts': parts, 'count': len(articles)}
    with open(os.path.join(OUT, 'index.json'), 'w', encoding='utf-8') as f:
        json.dump(index, f, ensure_ascii=False, separators=(',', ':'))
    with open(REPORT, 'w', encoding='utf-8') as f:
        f.write('\n'.join(report))
    print('\n'.join(report[:6]))
    print(f'报告: {REPORT}')


if __name__ == '__main__':
    main()
