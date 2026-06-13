# -*- coding: utf-8 -*-
"""
《印光大师文钞菁华录》（李净通编）原文本 + 白话本（曾琦云译注《说净土》）
→ 按数字编号配对成文白对照 实验模型 JSON

两个独立文件，靠条目数字编号（1…333）对齐：
  - 原文本 01菁华录-原文：10章（一…十）+ 子节（甲乙丙…）+ 333 条原文；
    末附「编者之言」「印造经像之功德」（白话本无，作纯原文附录）
  - 白话本 02菁华录-白话：9章 + 每章「导读」+ 数字编号白话（缺译 50/52/53/142）
分篇（同嘉言录体例）：有子节的章按子节分篇，无子节的章整章一篇；
章导读 → 该章首篇提要。缺译条按"有原文无白话"收录（前端单侧呈现）。
铁律：只配对、不改字；原文本/白话本各自逐字对账必须一致。

输出：build/jh/{index.json, articles/}
"""
import json
import os
import re

import docx

DIR = '/Users/bincai/Downloads/印光法师文钞word/嘉言录菁华录等'
SRC_O = os.path.join(DIR, '01菁华录-原文20201227.docx')
SRC_T = os.path.join(DIR, '02菁华录-白话20211227.docx')
PROJ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(PROJ, 'build', 'jh')
REPORT = os.path.join(PROJ, 'build', 'reports', 'report_jh.txt')

RE_CHAP = re.compile(r'^([一二三四五六七八九十]+)、([^，。、]{2,12})$')
RE_SUB = re.compile(r'^([甲乙丙丁戊己庚辛壬癸])、([^，。、]{2,12})$')
RE_O_NUM = re.compile(r'^(\d+)[\.\s、]?(\D.*)$', re.S)   # 原文本：1大矣哉 / 4.念佛
RE_T_NUM = re.compile(r'^(\d+)[、.]\s*(.*)$', re.S)        # 白话本：1、真是…


def sq(s):
    return re.sub(r'\s+', '', s)


def parse_orig():
    """原文本 → (条目 {n: (raw, text, chap, sub)}, 章节序, 附录段)"""
    d = docx.Document(SRC_O)
    paras = [p.text.strip() for p in d.paragraphs]
    # 正文起点：第二次出现「一、赞净土超胜」
    starts = [i for i, t in enumerate(paras) if t.startswith('一、赞净土超胜')]
    body = starts[1] if len(starts) > 1 else 0
    # 附录起点：「编者之言」
    appendix_start = next((i for i in range(body, len(paras))
                           if paras[i].startswith('编者之言')), len(paras))

    items = {}
    order = []          # [(类型, 文本)] 章/节顺序，建目录
    content = []        # 非标题内容行（对账流，独立于 items）
    chap = sub = ''
    for i in range(body, appendix_start):
        t = paras[i]
        if not t or 'PAGE' in t:
            continue
        mc = RE_CHAP.match(t)
        if mc and not RE_O_NUM.match(t):
            chap, sub = mc.group(0), ''
            order.append(('chap', chap))
            continue
        ms = RE_SUB.match(t)
        if ms:
            sub = ms.group(0)
            order.append(('sub', sub, chap))
            continue
        content.append(t)
        mn = RE_O_NUM.match(t)
        if mn and len(t) > 16:
            n = int(mn.group(1))
            items[n] = {'raw': t, 'text': mn.group(2).strip(), 'chap': chap, 'sub': sub}
        elif items:   # 条目续行（罕见）
            last = max(items)
            items[last]['text'] += '\n' + t
            items[last]['raw'] += '\n' + t

    appendix = [paras[i] for i in range(appendix_start, len(paras))
                if paras[i] and 'PAGE' not in paras[i]]
    return items, order, appendix, content


def parse_trans():
    """白话本 → (条目 {n: (raw, text)}, 章导读 {章名: 导读文})"""
    d = docx.Document(SRC_T)
    paras = [p.text.strip() for p in d.paragraphs]
    starts = [i for i, t in enumerate(paras) if t.startswith('第一章')]
    body = starts[-1] if len(starts) > 1 else starts[0]

    items = {}
    intros = {}          # 章序号(int) → 导读
    raws = []            # 对账用文档流（去章节标题/导读标题）
    chap_no = 0
    expect_intro = False
    for i in range(body, len(paras)):
        t = paras[i]
        if not t or 'PAGE' in t:
            continue
        if re.match(r'^第[一二三四五六七八九十]+章', t):
            chap_no += 1
            continue
        if t == '导读':
            expect_intro = True
            continue
        mn = RE_T_NUM.match(t)
        if mn and len(t) > 12:
            n = int(mn.group(1))
            items[n] = {'raw': t, 'text': mn.group(2).strip()}
            expect_intro = False
            raws.append(t)
            continue
        if expect_intro:
            intros.setdefault(chap_no, [])
            intros[chap_no].append(t)
            raws.append(t)
        else:
            # 条目续行（白话长条跨段）
            if items:
                last = max(items)
                items[last]['text'] += '\n' + t
                items[last]['raw'] += '\n' + t
                raws.append(t)
    intros = {k: '\n'.join(v) for k, v in intros.items()}
    return items, intros, raws


def main():
    o_items, order, appendix, o_content = parse_orig()
    t_items, intros, t_raws = parse_trans()
    report = ['《印光大师文钞菁华录》文白对照配对',
              f'原文本条目: {len(o_items)}（编号 {min(o_items)}-{max(o_items)}）',
              f'白话本条目: {len(t_items)}']
    miss_t = [n for n in o_items if n not in t_items]
    miss_o = [n for n in t_items if n not in o_items]
    report.append(f'有原文无白话（{len(miss_t)}）: {miss_t}')
    report.append(f'有白话无原文（{len(miss_o)}）: {miss_o}')

    # ---- 各自逐字对账（原文本：文档内容流 vs 解析条目流）----
    o_doc = sq(''.join(o_content))
    o_json = sq(''.join(o_items[n]['raw'] for n in sorted(o_items)))
    report.append(f"\n原文本对账: {'一致 ✓' if o_doc == o_json else f'⚠ 文档{len(o_doc)}/JSON{len(o_json)}'}")
    if o_doc != o_json:
        pos = next((k for k in range(min(len(o_doc), len(o_json))) if o_doc[k] != o_json[k]), 0)
        report.append(f'  首分歧@{pos}: 文档[…{o_doc[max(0,pos-8):pos+14]}] JSON[…{o_json[max(0,pos-8):pos+14]}]')
    t_doc = sq(''.join(t_raws))
    t_json_parts = []
    for cn in sorted(intros):
        t_json_parts.append(intros[cn])
    for n in sorted(t_items):
        t_json_parts.append(t_items[n]['raw'])
    # 导读与条目在文档中按章交错，单纯拼接顺序不同 → 仅比对字符集合长度+多重集
    from collections import Counter
    t_doc_c = Counter(t_doc)
    t_json_c = Counter(sq(''.join(t_json_parts)))
    report.append(f"白话本对账(字符多重集): {'一致 ✓' if t_doc_c == t_json_c else '⚠ 有差异'}")
    if t_doc_c != t_json_c:
        diff = (t_doc_c - t_json_c) + (t_json_c - t_doc_c)
        report.append(f'  差异字符: {dict(list(diff.items())[:10])}')

    # ---- 组织成文章（仿嘉言录分篇）----
    # 章 → 子节列表；判断章是否有子节
    chap_subs = {}
    for it in order:
        if it[0] == 'chap':
            chap_subs.setdefault(it[1], [])
        elif it[0] == 'sub':
            chap_subs.setdefault(it[2], []).append(it[1])

    chap_no_map = {}   # 章名 → 序号(int)，取导读
    for idx, (typ, *rest) in enumerate([o for o in order if o[0] == 'chap']):
        chap_no_map[rest[0]] = idx + 1

    articles = []
    seq = 0

    def items_of(chap, sub):
        return [n for n in sorted(o_items)
                if o_items[n]['chap'] == chap and o_items[n]['sub'] == sub]

    used_intro = set()
    for typ, *rest in order:
        if typ != 'chap':
            continue
        chap = rest[0]
        subs = chap_subs.get(chap, [])
        cn = chap_no_map.get(chap)
        if subs:
            # 章内可能有"无子节前言条目"(sub='')，并入第一子节前另立
            lead = items_of(chap, '')
            blocks = ([('', lead)] if lead else []) + [(s, items_of(chap, s)) for s in subs]
        else:
            blocks = [('', items_of(chap, ''))]
        first = True
        for sub, nums in blocks:
            if not nums:
                continue
            seq += 1
            title = sub if sub else chap
            segs = []
            for n in nums:
                seg = {'os': [o_items[n]['text']], 'ts': []}
                if n in t_items:
                    seg['ts'] = [t_items[n]['text']]
                segs.append(seg)
            art = {
                'id': f'{seq:03d}', 'title': title, 'translator': '',
                'summary': intros.get(cn, '') if first else '',
                'part': chap, 'segments': segs,
            }
            if first and cn in intros:
                used_intro.add(cn)
            first = False
            articles.append(art)

    # 附录：编者之言 + 印造经像之功德（纯原文）
    if appendix:
        seq += 1
        articles.append({
            'id': f'{seq:03d}', 'title': '编者之言 · 印造经像之功德', 'translator': '',
            'summary': '', 'part': '附录',
            'segments': [{'os': [x], 'ts': []} for x in appendix],
        })

    # ---- 输出 ----
    os.makedirs(os.path.join(OUT, 'articles'), exist_ok=True)
    for fn in os.listdir(os.path.join(OUT, 'articles')):
        os.remove(os.path.join(OUT, 'articles', fn))
    parts = []
    paired = 0
    for a in articles:
        with open(os.path.join(OUT, 'articles', a['id'] + '.json'), 'w', encoding='utf-8') as f:
            json.dump(a, f, ensure_ascii=False, separators=(',', ':'))
        if not parts or parts[-1]['title'] != a['part']:
            parts.append({'title': a['part'], 'articles': []})
        parts[-1]['articles'].append({'id': a['id'], 'title': a['title'], 'summary': a['summary']})
        paired += sum(1 for s in a['segments'] if s['ts'])
    index = {'id': 'jh', 'title': '印光大师文钞菁华录（文白对照）',
             'parts': parts, 'count': len(articles)}
    with open(os.path.join(OUT, 'index.json'), 'w', encoding='utf-8') as f:
        json.dump(index, f, ensure_ascii=False, separators=(',', ':'))
    report.append(f'\n文章数: {len(articles)}，配对条目: {paired}/{len(o_items)}')
    report.append('--- 每篇 ---')
    for a in articles:
        nt = sum(1 for s in a['segments'] if s['ts'])
        report.append(f"{a['id']} {a['title'][:16]:<18} 条目{len(a['segments']):>3} 配对{nt:>3}  [{a['part'][:10]}]"
                      + ('  导读' if a['summary'] else ''))
    with open(REPORT, 'w', encoding='utf-8') as f:
        f.write('\n'.join(report))
    print('\n'.join(report[:8]))
    print(f'报告: {REPORT}')


if __name__ == '__main__':
    main()
